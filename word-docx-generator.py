# Word-docx-Generator
# ======================================================================================
# Word-docx-Generator is a Python utility that uses a word docx file as template to 
# fill in content and batch generate documents. Batch content can be input from an excel 
# spreadsheet or csv file.

#%%

from io import StringIO
import shutil
import pandas as pd
import docx

template_path = "Template/Template.docx"
batch_content_path = "Content/Contract_Contents.xlsx"
filename_prefix = "15TG-001-CONTRACT_SPECS_"
filename_suffix = ".docx"


def main():
    batch_content = pd.read_excel(batch_content_path)

    for content in batch_content[['Section', 'Title']].values:
        if pd.notna(content[0]): section_number = content[0].strip()
        if pd.notna(content[1]): section_title = content[1].strip()

        if section_number.replace(" ","").isnumeric():
            destination_path = "Output/" + filename_prefix + section_number + filename_suffix
            shutil.copyfile(template_path, destination_path)
            print ('Copied file: ' + destination_path)

            template_fields_values = [['{XX XX XX}', section_number],
                ['{Section Title}', section_title],
                ['{SECTION TITLE}', section_title.upper()]]


            word_doc_replace_text(destination_path, template_fields_values)


def word_doc_replace_text(filename, template_field_values):

    document: docx.Document
    document = docx.Document(filename)

    # Search body
    for paragraph in document.paragraphs:
        paragraph_replace_text(paragraph, template_field_values)

    # Search header/footer sections
    for section in document.sections:
        for header in [section.header, section.first_page_header, 
            section.even_page_header]:
            for paragraph in header.paragraphs:
                paragraph_replace_text(paragraph, template_field_values)

        for footer in [section.footer, section.first_page_footer, 
            section.even_page_footer]:        
            for paragraph in footer.paragraphs:
                paragraph_replace_text(paragraph, template_field_values)

 
    document.save(filename)
    print('Finished processing: ' + filename)


def paragraph_replace_text(paragraph, template_field_values):
    for find_replace_pair in template_field_values:
        find_text = find_replace_pair[0]
        replace_text = find_replace_pair[1]
        if find_text in paragraph.text:
            paragraph.text = paragraph.text.replace(find_text, replace_text)


if __name__ == '__main__':
    main()
    print('Exiting ...')

# %%
